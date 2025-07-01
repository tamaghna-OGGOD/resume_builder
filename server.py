"""
server.py

FastAPI wrapper around the résumé-builder pipeline in script.py.
Provides:
  • GET  /           → HTML form with Model Type dropdown sourced from config.yml
  • POST /response   → Runs pipeline; shows JSON + Download Word link
"""

from __future__ import annotations
import io, shutil, tempfile, traceback, sys, base64
from pathlib import Path

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import HTMLResponse
from docx import Document

# Dynamically load the pipeline module (script.py)
import importlib.util as _import_util
ROOT = Path(__file__).parent
_spec = _import_util.spec_from_file_location("pipeline", ROOT / "script.py")
pipeline = _import_util.module_from_spec(_spec)
_spec.loader.exec_module(pipeline)  # type: ignore

app = FastAPI(title="ATS Résumé Builder")


def _run_pipeline(
    resume_file: UploadFile, jd_text: str | None, model_type: str, style: str
) -> dict:
    """
    Save PDF, invoke pipeline with JD and selected model, return result dict.
    """
    if resume_file.content_type != "application/pdf":
        raise HTTPException(400, "Résumé must be a PDF file")

    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = Path(tmpdir) / resume_file.filename
        with open(pdf_path, "wb") as f:
            shutil.copyfileobj(resume_file.file, f)

        try:
            return pipeline.run_pipeline(
                resume_path=str(pdf_path),
                job_description_text=jd_text,
                poppler_path=pipeline.CONFIG.get("poppler_path"),
                model_type=model_type,
                dpi=pipeline.CONFIG.get("dpi", 300),
                lang=pipeline.CONFIG.get("lang", "eng"),
                resume_style=style,
            )
        except Exception as exc:
            traceback.print_exc(file=sys.stderr)
            raise HTTPException(500, str(exc)) from exc

@app.get("/", response_class=HTMLResponse)
def form_view() -> HTMLResponse:
    """
    Serve a split UI:
    - Left: paste JD
    - Right: upload résumé
    - Dropdown: models from config.yml
    - Single Generate button (POST /response)
    """
    # Read model types from config.yml
    model_types = pipeline.CONFIG.get("model_types", {})
    options_html = "".join(
        f'<option value="{name}">{name}</option>' for name in model_types.keys()
    )

    # Read resume styles from config.yml
    styles = pipeline.CONFIG.get("resume_styles", [])
    style_options_html = "".join(
        f'<option value="{s}">{s.title()}</option>' for s in styles
    )

    html = f"""
    <!DOCTYPE html>
    <html lang="en">
      <head>
        <meta charset="UTF-8">
        <title>ATS Résumé Builder</title>
        <style>
          body {{ font-family: sans-serif; margin: 20px; }}
          .container {{ display: flex; gap: 20px; }}
          .panel {{ flex: 1; }}
          textarea {{ width: 100%; height: 200px; }}
          input[type=file] {{ width: 100%; }}
          select, button {{ margin-top: 10px; width: 100%; padding: 8px; }}
          label {{ display: block; margin-top: 10px; }}
        </style>
      </head>
      <body>
        <h1>ATS Résumé Builder</h1>
        <form method="post" action="/response" enctype="multipart/form-data">
          <div class="container">
            <div class="panel">
              <label for="model_type">Model Type:</label>
              <select name="model_type" id="model_type">
                {options_html}
              </select>
            </div>
            <div class="panel">
              <label for="style">Resume Style:</label>
              <select name="style" id="style">
                {style_options_html}
              </select>
            </div>
          </div>

          <div class="container">
            <div class="panel">
              <label>Job Description (paste here):</label>
              <textarea name="job_description" placeholder="Paste job description…"></textarea>
            </div>
            <div class="panel">
              <label>Upload Résumé (PDF):</label>
              <input type="file" name="resume" accept="application/pdf" required />
            </div>
          </div>

          <button type="submit">Generate</button>
        </form>
      </body>
    </html>
    """
    return HTMLResponse(html)

@app.post("/response", response_class=HTMLResponse)
async def response_endpoint(
    resume: UploadFile = File(...),
    job_description: str | None = Form(None),
    model_type: str = Form(...),
    style: str = Form(...),
):
    # Run pipeline
    result = _run_pipeline(resume, job_description, model_type, style)

    # Generate Word .docx for only rewritten résumé
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("Rewritten Résumé", level=1)
    doc.add_paragraph(result.get("rewritten_resume", "(none)"))
    doc.save(buf)
    buf.seek(0)
    b64 = base64.b64encode(buf.read()).decode()

    import json
    pretty = json.dumps(result, indent=2)

    html = f"""
    <!DOCTYPE html>
    <html lang="en">
      <head>
        <meta charset="UTF-8">
        <title>Response</title>
        <style>
          body {{ font-family: monospace; white-space: pre-wrap; margin: 20px; }}
          .download {{ display: block; margin-top: 20px; font-size: 1.2em; }}
        </style>
      </head>
      <body>
        <h1>Pipeline Output (JSON)</h1>
        <pre>{pretty}</pre>
        <a class="download" href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="revised_resume.docx">⬇️ Download Word (.docx)</a>
      </body>
    </html>
    """
    return HTMLResponse(html)
